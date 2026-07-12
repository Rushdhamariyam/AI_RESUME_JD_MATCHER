import streamlit as st
import spacy
import re
import io
import openai
from datetime import datetime
from docx import Document
from text_extraction import extract_text_from_pdf, clean_text
from matching import (
    calculate_similarity,
    extract_predefined_skills,
    extract_ner_skills,
    get_missing_skills
)

# -----------------------------------------------------------------------------
# 1. PAGE SETUP & CONFIGURATION
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title="AI Resume JD Matcher",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# -----------------------------------------------------------------------------
# 2. DESIGN SYSTEM
# -----------------------------------------------------------------------------
# Modern, professional design system with a clean, tech-forward aesthetic.
#
# Primary (Indigo):  #4F46E5
# Background:        #F8FAFC
# Card Background:   #FFFFFF
# Text Primary:      #0F172A
# Text Muted:        #64748B
# Border Color:      #E2E8F0
st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&family=Inter:wght@400;500;600;700&display=swap');

        :root {
            --primary: #4F46E5;
            --primary-hover: #4338CA;
            --bg-main: #F8FAFC;
            --bg-card: #FFFFFF;
            --text-main: #0F172A;
            --text-muted: #64748B;
            --border-color: #E2E8F0;
        }

        html, body, [class*="css"], .stApp {
            font-family: 'Inter', -apple-system, sans-serif !important;
            background-color: var(--bg-main) !important;
            color: var(--text-main) !important;
        }

        .block-container {
            padding-top: 2rem !important;
            padding-bottom: 4rem !important;
            padding-left: 2rem !important;
            padding-right: 2rem !important;
            max-width: 1200px !important;
        }

        /* ---------- Header Section ---------- */
        .app-header {
            text-align: center;
            margin-bottom: 2rem;
            padding: 1rem 0;
        }
        .app-title {
            font-family: 'Plus Jakarta Sans', sans-serif;
            font-weight: 800;
            font-size: 2.6rem;
            color: #1E293B;
            margin: 0;
            letter-spacing: -0.025em;
            background: linear-gradient(135deg, #4F46E5 0%, #06B6D4 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }
        .app-subtitle {
            font-family: 'Inter', sans-serif;
            font-size: 1.05rem;
            color: var(--text-muted);
            margin-top: 0.5rem;
            font-weight: 400;
        }

        hr.subtle-divider {
            border: none;
            border-top: 1px solid var(--border-color);
            margin: 1.5rem 0 2rem 0;
        }

        /* ---------- Panels & Columns ---------- */
        div[data-testid="column"] {
            background-color: var(--bg-card) !important;
            padding: 2rem !important;
            border-radius: 12px !important;
            border: 1px solid var(--border-color) !important;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05), 0 2px 4px -2px rgba(0, 0, 0, 0.05) !important;
        }

        .section-header {
            font-family: 'Plus Jakarta Sans', sans-serif;
            font-size: 1.25rem;
            font-weight: 700;
            color: #1E293B;
            border-bottom: 1px solid var(--border-color);
            padding-bottom: 12px;
            margin-bottom: 1.5rem;
        }

        /* ---------- Inputs ---------- */
        div[data-testid="stFileUploader"] {
            border: 1.5px dashed #CBD5E1 !important;
            border-radius: 8px !important;
            background-color: #F8FAFC !important;
            padding: 8px !important;
        }
        div[data-testid="stFileUploader"] section {
            background-color: transparent !important;
            border: none !important;
        }
        div[data-testid="stTextArea"] textarea {
            border-radius: 8px !important;
            border: 1px solid var(--border-color) !important;
            background-color: #FFFFFF !important;
            font-family: 'Inter', sans-serif !important;
            font-size: 14px !important;
            color: #334155 !important;
        }
        div[data-testid="stTextArea"] textarea:focus {
            border-color: var(--primary) !important;
            box-shadow: 0 0 0 2px rgba(99, 102, 241, 0.15) !important;
        }
        label, .stMarkdown p { 
            color: var(--text-main) !important; 
        }

        /* ---------- Action button ---------- */
        button[data-testid="stBaseButton-primary"] {
            background-color: var(--primary) !important;
            color: #FFFFFF !important;
            border: none !important;
            border-radius: 8px !important;
            padding: 0.6rem 1.5rem !important;
            font-weight: 600 !important;
            font-size: 14px !important;
            width: 100% !important;
            margin-top: 0.6rem;
            transition: all 0.2s ease-in-out !important;
            box-shadow: 0 1px 2px 0 rgba(0, 0, 0, 0.05) !important;
        }
        button[data-testid="stBaseButton-primary"]:hover {
            background-color: var(--primary-hover) !important;
            transform: translateY(-1px);
            box-shadow: 0 4px 6px -1px rgba(79, 70, 229, 0.2), 0 2px 4px -2px rgba(79, 70, 229, 0.2) !important;
        }

        /* ---------- Modern Scorecard ---------- */
        .score-card {
            border-radius: 12px;
            padding: 1.5rem;
            text-align: center;
            border: 1px solid var(--border-color);
            margin-bottom: 1.5rem;
            box-shadow: inset 0 2px 4px 0 rgba(0, 0, 0, 0.02);
        }
        .score-value {
            font-size: 3rem;
            font-weight: 800;
            line-height: 1;
            margin-bottom: 0.25rem;
            font-family: 'Plus Jakarta Sans', sans-serif;
        }
        .score-verdict {
            font-size: 1rem;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            margin-bottom: 0.25rem;
        }
        .score-label {
            font-size: 0.75rem;
            color: var(--text-muted);
            text-transform: uppercase;
            letter-spacing: 0.05em;
        }

        /* ---------- Skill Tags ---------- */
        .tag {
            display: inline-block;
            font-family: 'Inter', sans-serif;
            font-size: 12px;
            font-weight: 500;
            padding: 4px 10px;
            margin: 4px 6px 4px 0;
            border-radius: 6px;
            border-width: 1px;
            border-style: solid;
            transition: all 0.15s ease;
        }
        .tag:hover {
            transform: translateY(-1px);
        }
        .tag-flag {
            color: #E11D48;
            border-color: #FFE4E6;
            background: #FFF1F2;
        }
        .tag-flag::before { content: "\2717  "; }
        .tag-clear {
            color: #0D9488;
            border-color: #CCFBF1;
            background: #F0FDFA;
        }
        .tag-clear::before { content: "\2713  "; }
        .tag-neutral {
            color: #475569;
            border-color: #E2E8F0;
            background: #F1F5F9;
        }

        .section-label {
            font-family: 'Inter', sans-serif;
            font-size: 13px;
            font-weight: 600;
            color: #1E293B;
            margin: 1.5rem 0 0.25rem 0;
        }
        .section-caption {
            font-family: 'Inter', sans-serif;
            font-size: 12px;
            color: var(--text-muted);
            margin-bottom: 0.75rem;
        }

        /* ---------- Analyst's Note (Alert Box) ---------- */
        .info-note {
            background: #EEF2FF;
            border: 1px solid #E0E7FF;
            border-radius: 8px;
            padding: 1.25rem;
            margin-top: 1.8rem;
            font-family: 'Inter', sans-serif;
        }
        .info-header {
            font-weight: 700;
            font-size: 12px;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            color: #4F46E5;
            margin-bottom: 0.5rem;
        }
        .info-body {
            font-size: 12.5px;
            line-height: 1.6;
            color: #3730A3;
        }

        /* ---------- Empty State ---------- */
        .empty-state {
            border: 1.5px dashed var(--border-color);
            border-radius: 8px;
            padding: 3rem 1.5rem;
            text-align: center;
            color: var(--text-muted);
            font-family: 'Inter', sans-serif;
            font-size: 14px;
            line-height: 1.6;
        }
        .empty-state-title {
            font-weight: 600;
            font-size: 15px;
            color: #475569;
            margin-bottom: 0.5rem;
        }

        /* ---------- Improvement & Success Cards ---------- */
        .improvement-card {
            background: #FFFBEB;
            border: 1px solid #FEF3C7;
            border-radius: 12px;
            padding: 1.5rem;
            margin-bottom: 1.5rem;
        }
        .improvement-title {
            font-family: 'Plus Jakarta Sans', sans-serif;
            font-weight: 700;
            font-size: 1.1rem;
            color: #B45309;
            margin-bottom: 0.75rem;
            display: flex;
            align-items: center;
            gap: 8px;
        }
        .improvement-list {
            margin: 0;
            padding-left: 1.25rem;
            font-family: 'Inter', sans-serif;
            font-size: 13.5px;
            color: #78350F;
            line-height: 1.6;
        }
        .improvement-list li {
            margin-bottom: 0.5rem;
        }
        .improvement-list li:last-child {
            margin-bottom: 0;
        }
        .success-card {
            background: #F0FDFA;
            border: 1px solid #CCFBF1;
            border-radius: 12px;
            padding: 1.5rem;
            margin-bottom: 1.5rem;
            text-align: center;
        }
        .success-title {
            font-family: 'Plus Jakarta Sans', sans-serif;
            font-weight: 700;
            font-size: 1.1rem;
            color: #0F766E;
            margin-bottom: 0.25rem;
        }
        .success-text {
            font-family: 'Inter', sans-serif;
            font-size: 13.5px;
            color: #115E59;
        }

        /* ---------- Outlined Secondary Button ---------- */
        button[data-testid="stBaseButton-secondary"] {
            background-color: transparent !important;
            color: var(--primary) !important;
            border: 1px solid var(--primary) !important;
            border-radius: 8px !important;
            padding: 0.6rem 1.5rem !important;
            font-weight: 600 !important;
            font-size: 14px !important;
            width: 100% !important;
            margin-top: 0.6rem;
            transition: all 0.2s ease-in-out !important;
        }
        button[data-testid="stBaseButton-secondary"]:hover {
            background-color: rgba(79, 70, 229, 0.05) !important;
            border-color: var(--primary-hover) !important;
            color: var(--primary-hover) !important;
        }
    </style>
""", unsafe_allow_html=True)

# -----------------------------------------------------------------------------
# 3. CACHING & MODEL LOADING
# -----------------------------------------------------------------------------
@st.cache_resource
def load_nlp_models():
    """Downloads and caches the spaCy model and SBERT weights."""
    try:
        nlp = spacy.load("en_core_web_sm")
    except IOError:
        import subprocess
        import sys
        subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"], stdout=subprocess.DEVNULL)
        nlp = spacy.load("en_core_web_sm")

    from matching import get_sentence_transformer_model
    model = get_sentence_transformer_model('all-MiniLM-L6-v2')
    return nlp, model

with st.spinner("Loading language models..."):
    nlp, sbert_model = load_nlp_models()

# -----------------------------------------------------------------------------
# 4. UI RENDERING UTILITIES
# -----------------------------------------------------------------------------
def render_tags(skills, kind="neutral"):
    """Renders a set of skills as modern colored tags."""
    if not skills:
        st.markdown('<span style="font-size:12.5px; color:#64748B;">None identified.</span>', unsafe_allow_html=True)
        return
    css_class = {"flag": "tag-flag", "clear": "tag-clear", "neutral": "tag-neutral"}[kind]
    html = "".join([f'<span class="tag {css_class}">{s}</span>' for s in sorted(skills)])
    st.markdown(html, unsafe_allow_html=True)


def get_actionable_suggestion(item):
    """Generates a short, actionable suggestion for adding a missing skill."""
    item_lower = item.lower()
    
    # Certifications check
    if "cert" in item_lower or any(c in item_lower for c in ["cka", "pmp", "csm", "cissp", "scrummaster"]):
        return f"Include a certification in <strong>{item}</strong> if applicable, or highlight any corresponding training."
    
    # Database check
    databases = ["sql", "postgres", "mysql", "mongodb", "redis", "dynamodb", "neo4j", "sqlite", "oracle", "cassandra", "snowflake", "bigquery"]
    if any(db in item_lower for db in databases):
        return f"Specify database experience with <strong>{item}</strong>, including query optimization, schema design, or data warehousing."
        
    # Languages check
    languages = ["python", "julia", "scala", "java", "c++", "c#", "javascript", "matlab", "go", "bash", "rust", "sas", "ruby", "php"]
    if any(lang == item_lower for lang in languages):
        return f"Add a bullet point mentioning experience with <strong>{item}</strong> scripting or software development."
        
    # Cloud check
    cloud = ["aws", "azure", "gcp", "amazon web services", "google cloud"]
    if any(cl in item_lower for cl in cloud):
        return f"Highlight cloud infrastructure deployment, provisioning, or integration using <strong>{item}</strong>."

    # DevOps/CI-CD/Tools check
    devops = ["docker", "kubernetes", "git", "ci/cd", "jenkins", "terraform", "ansible", "airflow"]
    if any(do in item_lower for do in devops):
        return f"Mention workflows, pipelines, or version control tools involving <strong>{item}</strong>."

    # Concepts & Methodologies
    methodology = ["machine learning", "deep learning", "nlp", "computer vision", "statistics", "probability", "agile", "scrum", "regression", "classification", "clustering", "ab testing", "a/b testing"]
    if any(m in item_lower for m in methodology):
        return f"Showcase project experience or theoretical understanding of <strong>{item}</strong> methodologies."
        
    return f"Add a bullet point mentioning experience with <strong>{item}</strong>."


def render_verdict_stamp(score):
    """Renders the compatibility score as a clean, modern card."""
    if score >= 80:
        verdict, color, bg_color, border_color = "Strong Match", "#0D9488", "#F0FDFA", "#CCFBF1"
    elif score >= 50:
        verdict, color, bg_color, border_color = "Partial Match", "#D97706", "#FFFBEB", "#FEF3C7"
    else:
        verdict, color, bg_color, border_color = "Weak Match", "#E11D48", "#FFF1F2", "#FFE4E6"

    st.markdown(f"""
        <div class="score-card" style="background-color: {bg_color}; border-color: {border_color};">
            <div class="score-value" style="color: {color};">{score:.0f}%</div>
            <div class="score-verdict" style="color: {color};">{verdict}</div>
            <div class="score-label">Compatibility Index</div>
        </div>
    """, unsafe_allow_html=True)


def generate_tailored_resume_llm(original_resume, jd_text, matched_skills, missing_skills):
    """
    Attempts to call OpenAI API to generate a tailored resume.
    """
    try:
        client = openai.OpenAI()
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system", 
                    "content": (
                        "You are an expert ATS optimization resume assistant. Your task is to rewrite a resume "
                        "to align with a job description. Ensure that you never invent new job roles, projects, "
                        "achievements, or certifications. Only rephrase existing bullet points to use the "
                        "matching vocabulary of the JD and reorder sections to highlight matches."
                    )
                },
                {
                    "role": "user",
                    "content": f"""
Please optimize and rewrite the following candidate resume for the target job description.

REAL EXPERIENCE CONSTRAINTS:
- Do not fabricate new jobs, degrees, certifications, or projects.
- Keep all original metrics, roles, and dates identical.
- Rephrase descriptions and bullet points to naturally incorporate job keywords/skills where applicable.
- In a section at the top titled "[⚠️ UNADDRESSABLE REQUIREMENTS / GAP NOTICE]", list any key requirements from the JD (like years of experience, specific degrees, or technologies) that could not be addressed in the rewrite.

Match Context:
- Confirmed Matches: {', '.join(matched_skills)}
- Missing Skills: {', '.join(missing_skills)}

Original Resume:
{original_resume}

Job Description:
{jd_text}
"""
                }
            ],
            temperature=0.3
        )
        return response.choices[0].message.content, None
    except Exception as e:
        return None, str(e)


def generate_tailored_resume_fallback(original_resume, jd_text, matched_skills, missing_skills):
    """
    Local fallback rule-based rewriter when the LLM API is unavailable.
    """
    lines = original_resume.split('\n')
    tailored_lines = []
    injected_skills = False
    
    for line in lines:
        cleaned_line = line.strip()
        if re.search(r'(?i)\b(?:technical\s+)?skills\b', cleaned_line) and not injected_skills:
            tailored_lines.append(line)
            if missing_skills:
                suggestions = ", ".join(list(missing_skills)[:5])
                tailored_lines.append(f"  - Key Technologies (ATS Matched): {', '.join(matched_skills)}")
                tailored_lines.append(f"  - Recommended Additions (Familiarity): {suggestions}")
                injected_skills = True
            continue
            
        rephrased = line
        for skill in matched_skills:
            pattern = r'(?i)\b' + re.escape(skill) + r'\b'
            rephrased = re.sub(pattern, skill, rephrased)
            
        tailored_lines.append(rephrased)
        
    tailored_text = "\n".join(tailored_lines)
    
    unaddressed = list(missing_skills)[5:] if len(missing_skills) > 5 else []
    unaddressed_text = ""
    if unaddressed:
        unaddressed_text = "\n".join([f"- {item}" for item in unaddressed])
    else:
        unaddressed_text = "- None (All missing keywords were integrated as suggestions)"
        
    header = f"""[⚠️ UNADDRESSABLE REQUIREMENTS / GAP NOTICE]
The following key requirements from the Job Description could not be integrated
automatically based on your resume profile. Manually verify if you possess these:
{unaddressed_text}

====================================================================
"""
    return header + "\n" + tailored_text


def export_to_docx(text):
    """
    Generates a docx binary file from resume text.
    """
    doc = Document()
    lines = text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
            
        if stripped.startswith("===") or stripped.startswith("---"):
            doc.add_paragraph("__________________________________________________")
        elif stripped.startswith("[⚠️") or (stripped.isupper() and len(stripped) < 50):
            doc.add_heading(stripped, level=2)
        elif stripped.startswith("-") or stripped.startswith("*"):
            doc.add_paragraph(stripped.lstrip("-* ").strip(), style='List Bullet')
        else:
            doc.add_paragraph(stripped)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# -----------------------------------------------------------------------------
# 5. APPLICATION HEADER
# -----------------------------------------------------------------------------
st.markdown("""
    <div class="app-header">
        <div class="app-title">AI Resume JD Matcher</div>
        <div class="app-subtitle">Semantic compatibility analysis using sentence embeddings &amp; entity extraction</div>
    </div>
""", unsafe_allow_html=True)

st.markdown('<hr class="subtle-divider">', unsafe_allow_html=True)

# -----------------------------------------------------------------------------
# 6. WORKSPACE LAYOUT
# -----------------------------------------------------------------------------
col_input, col_results = st.columns([1.1, 1], gap="large")

with col_input:
    st.markdown('<div class="section-header">Upload Documents</div>', unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "Resume (PDF)",
        type=["pdf"],
        help="Upload the candidate resume in PDF format."
    )

    st.markdown('<div style="margin-bottom: 1rem;"></div>', unsafe_allow_html=True)

    jd_input = st.text_area(
        "Target job description",
        height=250,
        placeholder="Paste the job description, requirements, and required skills here...",
        help="Paste the full text of the job posting."
    )

    st.markdown('<div style="margin-bottom: 1rem;"></div>', unsafe_allow_html=True)

    check_button = st.button("Check Match & Analyze", type="primary")

with col_results:
    st.markdown('<div class="section-header">Analysis Results</div>', unsafe_allow_html=True)

    # Initialize state defaults
    if "analyzed" not in st.session_state:
        st.session_state.analyzed = False
        st.session_state.match_score = 0.0
        st.session_state.all_resume_skills = set()
        st.session_state.all_jd_skills = set()
        st.session_state.missing_skills = []
        st.session_state.matched_skills = set()
        st.session_state.raw_resume_text = ""
        st.session_state.jd_input = ""
        st.session_state.tailored_resume = None

    if check_button:
        if not uploaded_file:
            st.error("Please upload a resume PDF before running the analysis.")
        elif not jd_input.strip():
            st.error("Please paste a job description before running the analysis.")
        else:
            with st.spinner("Analyzing matching score..."):
                raw_resume_text = extract_text_from_pdf(uploaded_file)
                clean_resume = clean_text(raw_resume_text)
                clean_jd = clean_text(jd_input)

                if not clean_resume:
                    st.error("Could not read text from this PDF. Make sure it isn't a scanned image.")
                else:
                    match_score = calculate_similarity(clean_resume, clean_jd, model=sbert_model)

                    resume_skills_pre = extract_predefined_skills(clean_resume)
                    resume_skills_ner = extract_ner_skills(raw_resume_text, resume_skills_pre, nlp)
                    all_resume_skills = resume_skills_pre.union(resume_skills_ner)

                    jd_skills_pre = extract_predefined_skills(clean_jd)
                    jd_skills_ner = extract_ner_skills(jd_input, jd_skills_pre, nlp)
                    all_jd_skills = jd_skills_pre.union(jd_skills_ner)

                    missing_skills = get_missing_skills(all_resume_skills, all_jd_skills)
                    matched_skills = all_jd_skills.difference(missing_skills)

                    # Update persistent state
                    st.session_state.match_score = match_score
                    st.session_state.all_resume_skills = all_resume_skills
                    st.session_state.all_jd_skills = all_jd_skills
                    st.session_state.missing_skills = missing_skills
                    st.session_state.matched_skills = matched_skills
                    st.session_state.raw_resume_text = raw_resume_text
                    st.session_state.jd_input = jd_input
                    st.session_state.analyzed = True
                    st.session_state.tailored_resume = None

    if st.session_state.analyzed:
        # Render using persistent state
        match_score = st.session_state.match_score
        all_resume_skills = st.session_state.all_resume_skills
        all_jd_skills = st.session_state.all_jd_skills
        missing_skills = st.session_state.missing_skills
        matched_skills = st.session_state.matched_skills
        raw_resume_text = st.session_state.raw_resume_text
        jd_input = st.session_state.jd_input

        render_verdict_stamp(match_score)

        # Compute frequency of each missing skill in the job description to rank them
        jd_text_lower = jd_input.lower()
        skill_frequencies = {}
        for skill in missing_skills:
            pattern = r'(?<![a-zA-Z0-9_])' + re.escape(skill.lower()) + r'(?![a-zA-Z0-9_])'
            count = len(re.findall(pattern, jd_text_lower))
            skill_frequencies[skill] = max(1, count)
        
        # Sort missing skills by frequency in descending order
        ranked_missing = sorted(missing_skills, key=lambda s: skill_frequencies[s], reverse=True)
        top_missing = ranked_missing[:5]

        if top_missing:
            li_items = "".join([f"<li>{get_actionable_suggestion(item)}</li>" for item in top_missing])
            st.markdown(f"""
                <div class="improvement-card">
                    <div class="improvement-title">💡 What to Add to Improve Your Match</div>
                    <ul class="improvement-list">
                        {li_items}
                    </ul>
                </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
                <div class="success-card">
                    <div class="success-title">✨ No Major Gaps Found</div>
                    <div class="success-text">No major gaps found — your resume covers the key requirements.</div>
                </div>
            """, unsafe_allow_html=True)

        # Tailor Resume Section
        st.markdown('<hr class="subtle-divider">', unsafe_allow_html=True)
        st.markdown('<div class="section-label">Tailored Resume</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-caption">Generate an ATS-optimized version of your resume aligned with the job description.</div>', unsafe_allow_html=True)
        
        generate_tailored_button = st.button("Generate Tailored Resume ✨", type="secondary", key="gen_tailored")

        if generate_tailored_button or st.session_state.tailored_resume is not None:
            if st.session_state.tailored_resume is None:
                with st.spinner("Rewriting/tailoring resume content..."):
                    tailored_res, err = generate_tailored_resume_llm(
                        raw_resume_text,
                        jd_input,
                        matched_skills,
                        missing_skills
                    )
                    
                    if tailored_res:
                        st.session_state.tailored_resume = tailored_res
                        st.success("Resume tailored successfully using OpenAI GPT!")
                    else:
                        # Fallback
                        st.session_state.tailored_resume = generate_tailored_resume_fallback(
                            raw_resume_text,
                            jd_input,
                            matched_skills,
                            missing_skills
                        )
                        st.info(f"OpenAI API quota exceeded ({err}). Tailored resume generated using offline fallback pipeline.")

            # Editable Text Area for tailored resume
            tailored_text_area = st.text_area(
                "Tailored Resume Content (Editable)",
                value=st.session_state.tailored_resume,
                height=350,
                key="tailored_text_area_field"
            )
            st.session_state.tailored_resume = tailored_text_area

            # Export/Download Buttons
            col_dl1, col_dl2 = st.columns(2)
            with col_dl1:
                docx_bytes = export_to_docx(st.session_state.tailored_resume)
                st.download_button(
                    label="📥 Download .docx Word File",
                    data=docx_bytes,
                    file_name="tailored_resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_docx"
                )
            with col_dl2:
                st.download_button(
                    label="📄 Download .txt Text File",
                    data=st.session_state.tailored_resume,
                    file_name="tailored_resume.txt",
                    mime="text/plain",
                    key="dl_txt"
                )

        st.markdown('<hr class="subtle-divider">', unsafe_allow_html=True)

        st.markdown('<div class="section-label">Missing from resume</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-caption">Required by the posting, not found in the resume.</div>', unsafe_allow_html=True)
        render_tags(missing_skills, kind="flag")

        st.markdown('<div class="section-label">Confirmed match</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-caption">Present in both the resume and the posting.</div>', unsafe_allow_html=True)
        render_tags(matched_skills, kind="clear")

        with st.expander("Full inventory \u2014 all skills detected in resume"):
            render_tags(all_resume_skills, kind="neutral")

        st.markdown(f"""
            <div class="info-note">
                <div class="info-header">Analyst's Note</div>
                <div class="info-body">
                    This score comes from a hybrid pipeline: exact-boundary keyword matching
                    catches named tools and languages, while an SBERT sentence embedding
                    (all-MiniLM-L6-v2) captures semantic overlap that keyword search alone
                    would miss &mdash; e.g. recognizing "predictive modelling" as related to
                    "regression" even without a literal word match.
                </div>
            </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
            <div class="empty-state">
                <div class="empty-state-title">Analysis Pending</div>
                Upload a resume and paste a job description under "Upload Documents",<br>
                then run the analysis to view results here.
            </div>
        """, unsafe_allow_html=True)
